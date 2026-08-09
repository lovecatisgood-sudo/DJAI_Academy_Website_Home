from __future__ import annotations

import csv
import io
import json
from pathlib import Path
from typing import Any


def timestamp(seconds: float, srt: bool = True) -> str:
    seconds = max(0.0, float(seconds))
    millis = int(round(seconds * 1000))
    hours, rem = divmod(millis, 3_600_000)
    minutes, rem = divmod(rem, 60_000)
    secs, ms = divmod(rem, 1000)
    sep = "," if srt else "."
    return f"{hours:02d}:{minutes:02d}:{secs:02d}{sep}{ms:03d}"


def speaker_prefix(segment: dict) -> str:
    speaker = segment.get("speaker")
    return f"{speaker}: " if speaker else ""


def render_txt(job: dict) -> bytes:
    result = job.get("result") or {}
    lines = []
    for seg in result.get("segments") or []:
        lines.append(f"[{timestamp(seg.get('start', 0), srt=False)}] {speaker_prefix(seg)}{seg.get('text','').strip()}")
    return ("\n".join(lines).strip() + "\n").encode("utf-8")


def render_srt(job: dict) -> bytes:
    out = []
    for i, seg in enumerate((job.get("result") or {}).get("segments") or [], 1):
        out.extend([
            str(i),
            f"{timestamp(seg.get('start',0))} --> {timestamp(seg.get('end',0))}",
            f"{speaker_prefix(seg)}{seg.get('text','').strip()}", ""
        ])
    return "\n".join(out).encode("utf-8")


def render_vtt(job: dict) -> bytes:
    out = ["WEBVTT", ""]
    for seg in (job.get("result") or {}).get("segments") or []:
        out.extend([
            f"{timestamp(seg.get('start',0), False)} --> {timestamp(seg.get('end',0), False)}",
            f"{speaker_prefix(seg)}{seg.get('text','').strip()}", ""
        ])
    return "\n".join(out).encode("utf-8")


def render_csv(job: dict) -> bytes:
    buf = io.StringIO()
    writer = csv.writer(buf)
    writer.writerow(["start", "end", "speaker", "text"])
    for seg in (job.get("result") or {}).get("segments") or []:
        writer.writerow([seg.get("start", 0), seg.get("end", 0), seg.get("speaker") or "", seg.get("text", "")])
    return buf.getvalue().encode("utf-8-sig")


def render_json(job: dict) -> bytes:
    payload = {
        "id": job.get("id"), "filename": job.get("filename"), "duration_seconds": job.get("duration_seconds"),
        "detected_language": job.get("detected_language"), "backend": job.get("backend"),
        "options": job.get("options"), "result": job.get("result"),
    }
    return json.dumps(payload, ensure_ascii=False, indent=2).encode("utf-8")


def render_docx(job: dict, output: Path) -> None:
    from docx import Document
    from docx.shared import Pt
    doc = Document()
    doc.add_heading("Transcript", 0)
    doc.add_paragraph(job.get("filename") or "Audio / video transcript")
    for seg in (job.get("result") or {}).get("segments") or []:
        p = doc.add_paragraph()
        r = p.add_run(f"{timestamp(seg.get('start',0), False)}  {seg.get('speaker') or ''}".strip())
        r.bold = True
        r.font.size = Pt(9)
        p.add_run("\n" + seg.get("text", "").strip())
    doc.save(output)


def render_pdf(job: dict, output: Path) -> None:
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_LEFT
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

    font = "Helvetica"
    for candidate in [
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
        Path("/usr/share/fonts/truetype/noto/NotoSansThai-Regular.ttf"),
    ]:
        if candidate.exists():
            try:
                pdfmetrics.registerFont(TTFont("SiameseUnicode", str(candidate)))
                font = "SiameseUnicode"
                break
            except Exception:
                pass
    styles = getSampleStyleSheet()
    title = ParagraphStyle("SiameseTitle", parent=styles["Title"], fontName=font, textColor=colors.HexColor("#126049"))
    meta = ParagraphStyle("SiameseMeta", parent=styles["Normal"], fontName=font, fontSize=9, textColor=colors.HexColor("#6b746f"), spaceAfter=2)
    body = ParagraphStyle("SiameseBody", parent=styles["BodyText"], fontName=font, fontSize=10.5, leading=15, alignment=TA_LEFT, spaceAfter=8)
    story = [Paragraph("Transcript", title), Paragraph(job.get("filename") or "Audio / video transcript", meta), Spacer(1, 5*mm)]
    for seg in (job.get("result") or {}).get("segments") or []:
        label = f"{timestamp(seg.get('start',0), False)}  {seg.get('speaker') or ''}".strip()
        story.append(Paragraph(label, meta))
        safe = (seg.get("text", "").strip().replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"))
        story.append(Paragraph(safe, body))
    doc = SimpleDocTemplate(str(output), pagesize=A4, rightMargin=18*mm, leftMargin=18*mm, topMargin=18*mm, bottomMargin=18*mm)
    doc.build(story)


MIME = {
    "txt": "text/plain; charset=utf-8", "srt": "application/x-subrip", "vtt": "text/vtt; charset=utf-8",
    "csv": "text/csv; charset=utf-8", "json": "application/json; charset=utf-8",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "pdf": "application/pdf",
}


def export_job(job: dict, fmt: str, output: Path) -> None:
    fmt = fmt.lower()
    if fmt == "txt": output.write_bytes(render_txt(job))
    elif fmt == "srt": output.write_bytes(render_srt(job))
    elif fmt == "vtt": output.write_bytes(render_vtt(job))
    elif fmt == "csv": output.write_bytes(render_csv(job))
    elif fmt == "json": output.write_bytes(render_json(job))
    elif fmt == "docx": render_docx(job, output)
    elif fmt == "pdf": render_pdf(job, output)
    else: raise ValueError(f"Unsupported export format: {fmt}")
